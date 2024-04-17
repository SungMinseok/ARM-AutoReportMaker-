import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def replace_text_with_format(paragraph, old_text, new_text):
    p = paragraph._element
    for run in p.iter(qn('w:r')):
        for text in run.iter(qn('w:t')):
            if old_text in text.text:
                # Save the formatting
                run_format = run.get_or_add_rPr()
                font = run_format.get_or_add_rFonts()
                size = run_format.get_or_add_sz()
                
                # Replace the text
                text.text = text.text.replace(old_text, new_text)
                print(text.text)
                
                # Apply the formatting to the new text
                # new_run = p.add_r()
                # new_run.append(run)
                # new_text_element = new_run.get_or_add_t()
                # new_text_element.text = new_text
                # font.set(qn('w:eastAsia'), font.eastAsia)
                # size.set(qn('w:val'), size.val)
                #break

def replace_values_in_word(input_docx, output_docx, replacements):
    #Load the original document
    doc = Document(input_docx)
    
    # Replace placeholders with values from the replacements dictionary
    for paragraph in doc.paragraphs:
        #print(paragraph.text)
        for key, value in replacements.items():
            try: 
                if key in paragraph.text:
                    #paragraph.text = paragraph.text.replace(key, value)
                    replace_text_with_format(paragraph, key, value)
            except : 
                print(f'{key}의 value 없음.')
                continue


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        try: 
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                replace_text_with_format(paragraph, key, value)

                        except : 
                            print(f'{key}의 value 없음.')
                            continue
        #print(table.rows)


    # Save the modified document
    doc.save(output_docx)
    # Load the text from the original document
    # text = docx2txt.process(input_docx)

    # # Replace placeholders with values from the replacements dictionary
    # for key, value in replacements.items():
    #     try:
    #         text = text.replace(key, value)
    #     except:
    #         print(f'{key}의 value 없음.')
    #         continue

    
    # # Save the modified text to a new document
    # with open(output_docx, 'w', encoding='utf-8') as file:
    #     file.write(text)

    # print(text)


def update_doc_by_info(input_name,output_name,info_file_name="",info_dict=None):
    #excel_filename = f'qa_info.csv'
    
    # Load Excel sheet
    replacements = {}

    if info_file_name != "" :
    
        df_info = pd.read_csv(info_file_name, index_col='Key')

        for key, value in df_info['Value'].items():
            replacements[key] = value

    elif info_dict != None :
        replacements = info_dict
    
    #input_docx = './양식/QA_결과 보고 문서.docx'
    #output_docx = os.path.join(output_path,f'QA_결과 보고_{country}_{project_name}_{date.strftime("%y%m%d")} {doc_type} QA 결과.docx')
    #output_docx = os.path.join(output_path,output_name)
    #os.startfile(input_name)
    replace_values_in_word(input_name, output_name, replacements)

    #os.startfile(output_docx)

if __name__  == "__main__" :
    # 호출 예시
    # country = 'TW'
    # date = pd.to_datetime('2023-08-16')
    # project_name = 'R2M'
    # doc_type = '업데이트'
    #update_qa_report(country, date, project_name, doc_type)
    name = 'QA_결과 보고_TW_R2M_231017 업데이트 QA 결과.docx'
    update_qa_report(name)
